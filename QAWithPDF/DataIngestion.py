from exception import customexception
from logger import logging
import sys
from llama_index.core import Document
import chardet
import PyPDF2
from docx import Document as DocxDocument
from io import BytesIO

class DataIngestion:
    def __init__(self):
        pass

    def initiate_data_ingestion(self, uploaded_files):
        """
        Initiates the data ingestion process by reading the contents of uploaded files.

        Args:
        - uploaded_files (list): List of Streamlit file objects.

        Returns:
        - docs (list): List of Document objects loaded from the uploaded files.
        """
        try:
            logging.info("Data ingestion initiated")
            docs = []
            for file in uploaded_files:
                file_name = getattr(file, 'name', '')
                if file_name:
                    # Use endswith to check the file extension
                    if file_name.endswith('.txt'):
                        content = self._handle_text_file(file)
                    elif file_name.endswith('.pdf'):
                        content = self._handle_pdf_file(file)
                    elif file_name.endswith('.docx'):
                        content = self._handle_docx_file(file)
                    else:
                        logging.warning(f"Unsupported file type: {file_name}. Skipping.")
                        continue

                    # Log the content (first 500 characters) for debugging
                    logging.info(f"Extracted content from {file_name}: {content[:500]}")

                    # Add the content to the docs list
                    docs.append(Document(text=content))  # Ensure 'text' argument is used
                    logging.info(f"Loaded document from {file_name}")
                else:
                    logging.warning("File has no name attribute, skipping.")

            if not docs:
                logging.warning("No documents were loaded from the uploaded files.")
            print(docs[0])

            return docs
        except Exception as e:
            logging.error("Error during data ingestion", exc_info=True)
            raise customexception(e, sys)

    def _handle_text_file(self, file):
        """Handles text files and returns their content."""
        file_bytes = file.read()
        result = chardet.detect(file_bytes)
        encoding = result['encoding'] if result['encoding'] else 'utf-8'
        return file_bytes.decode(encoding)

    def _handle_pdf_file(self, file):
        """Handles PDF files and extracts their text."""
        reader = PyPDF2.PdfReader(file)
        text = ''
        for page in reader.pages:
            text += page.extract_text() or ''
        if not text:
            logging.warning(f"No text extracted from PDF file {file.name}")
        return text

    def _handle_docx_file(self, file):
        """Handles DOCX files and extracts their text, returning as a Document."""
        doc = DocxDocument(BytesIO(file.read()))
        text = ' '.join([para.text for para in doc.paragraphs])
        if not text:
            logging.warning(f"No text extracted from DOCX file {file.name}")
        return text  # Return the content as text
